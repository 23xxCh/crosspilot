#!/usr/bin/env python3
"""下载图片+翻译+输出DOCX"""
import json, os, sys, time, glob, requests
from concurrent.futures import ThreadPoolExecutor, as_completed

if __package__ in {None, ""}:
    from _bootstrap import ensure_package_imports
    ensure_package_imports()

from scripts.model_provider import get_provider

OUT = r'E:\AI WORK\ebay -skill\图片和翻译文件'
IMG_DIR = os.path.join(OUT, '图片')
os.makedirs(IMG_DIR, exist_ok=True)

# 1. 读取输出
files = glob.glob(r'亚马逊表\跨境电商自动化回填表2_121050_586.json')
if not files:
    files = sorted(glob.glob(r'亚马逊表\跨境电商自动化回填表2_*.json'), key=os.path.getmtime)
    # 排除 doublecheck 和 metrics
    files = [f for f in files if 'doublecheck' not in f and 'metrics' not in f and 'r' not in os.path.basename(f).split('_')[-1][:2]]
    files = [max(files, key=os.path.getmtime)] if files else []

output_path = files[0]
print(f'Source: {os.path.basename(output_path)}')

with open(output_path, encoding='utf-8') as f:
    data = json.load(f)

fields = list(data.keys())
title_f, desc_f, img_f, var_f = fields[1], fields[2], fields[3], fields[4]
bullet_fs = [f'Bullet Point{i}' for i in range(1,6)]
kw_f = fields[10]
rows = len(data[title_f])

# 2. 收集所有图片URL并下载
all_urls = set()
url_map = {}  # row -> [main_url, extra_urls, var_urls]
for i in range(rows):
    imgs = data[img_f][i] if i < len(data[img_f]) else []
    vars_ = data[var_f][i] if i < len(data[var_f]) else []
    main = imgs[0] if imgs else ''
    extra = imgs[1:] if len(imgs) > 1 else []
    if main: all_urls.add(main)
    for u in extra: all_urls.add(u)
    for u in vars_: all_urls.add(u)
    url_map[i] = (main, extra, vars_)

print(f'Images to download: {len(all_urls)}')

# 下载
dl_lock = __import__('threading').Lock()
dl_done = dl_ok = 0

def download(url):
    global dl_done, dl_ok
    fname = os.path.join(IMG_DIR, url.split('/')[-1].split('?')[0])
    if not fname.endswith('.png'): fname += '.png'
    if os.path.exists(fname):
        with dl_lock: dl_done += 1; dl_ok += 1
        return fname
    try:
        r = requests.get(url, timeout=30, stream=True)
        if r.ok:
            with open(fname, 'wb') as f:
                for chunk in r.iter_content(8192): f.write(chunk)
            with dl_lock: dl_done += 1; dl_ok += 1
            return fname
    except: pass
    with dl_lock: dl_done += 1
    return ''

t0 = time.time()
with ThreadPoolExecutor(max_workers=10) as pool:
    futures = [pool.submit(download, u) for u in all_urls]
    for _ in as_completed(futures):
        if dl_done % 100 == 0: print(f'  Download: {dl_done}/{len(all_urls)}')
print(f'Download done: {dl_ok}/{len(all_urls)} in {time.time()-t0:.0f}s')

# 3. 翻译标题和描述到中文
print('\nTranslating...')
provider = get_provider()

# 批处理翻译
def batch_translate(texts, label):
    results = {}
    batch_size = 15
    batches = [texts[i:i+batch_size] for i in range(0, len(texts), batch_size)]
    tdone = [0]
    lock = __import__('threading').Lock()

    def process(batch):
        indexed = '\n'.join([f'[{j}] {t}' for j, t in enumerate(batch)])
        prompt = f'Translate these to Chinese. Return JSON array: [{{"index":N,"text":"..."}}]. Keep product specs intact.\n\n{indexed}'
        result = provider.call_text(prompt, max_tokens=4096)
        parsed = {}
        if result:
            try:
                items = json.loads(result)
                for item in items:
                    if isinstance(item, dict) and 'index' in item:
                        parsed[item['index']] = item.get('text', '')
            except: pass
        with lock: tdone[0] += 1
        return {batch[i]: parsed.get(i, '') for i in range(len(batch))}

    all_results = {}
    with ThreadPoolExecutor(max_workers=5) as pool:
        futures = {pool.submit(process, b): b for b in batches}
        for future in as_completed(futures):
            all_results.update(future.result())
            if tdone[0] % 5 == 0: print(f'  {label}: {tdone[0]}/{len(batches)} batches')
    return all_results

# 翻译标题
titles_to_translate = [data[title_f][i] for i in range(rows)]
zh_titles = batch_translate(titles_to_translate, 'Titles')

# 翻译描述
descs_to_translate = [data[desc_f][i] for i in range(rows)]
zh_descs = batch_translate(descs_to_translate, 'Descriptions')

# 4. 生成 DOCX
print('\nGenerating DOCX...')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Amazon 回填表 - 产品清单', 0)
doc.add_paragraph(f'共 {rows} 个产品 | 生成时间: {time.strftime("%Y-%m-%d %H:%M")}')

for i in range(rows):
    title_en = data[title_f][i]
    desc_en_raw = data[desc_f][i]
    title_zh = zh_titles.get(title_en, '')
    desc_zh = zh_descs.get(desc_en_raw, '')
    main_img, extra_imgs, var_imgs = url_map[i]

    # 产品标题
    doc.add_heading(f'#{i+1} {title_en}', level=2)
    if title_zh:
        doc.add_paragraph(f'[中文] {title_zh}').bold = True

    # 描述
    doc.add_heading('产品描述', level=3)
    doc.add_paragraph(desc_en_raw[:500])
    if desc_zh:
        p = doc.add_paragraph()
        run = p.add_run(f'[中文] {desc_zh[:500]}')
        run.font.color.rgb = None  # 黑色

    # Keywords
    kw = data[kw_f][i]
    if kw:
        doc.add_paragraph(f'关键词: {kw[:200]}')

    # Bullet Points
    bullets = [data[f][i] for f in bullet_fs]
    if any(bullets):
        doc.add_heading('Bullet Points', level=3)
        for b in bullets:
            if b: doc.add_paragraph(b[:300], style='List Bullet')

    # 图片
    doc.add_heading('产品图片', level=3)
    all_imgs = [main_img] + extra_imgs + var_imgs
    for j, url in enumerate(all_imgs[:10]):  # 最多10张
        if not url: continue
        fname = url.split('/')[-1].split('?')[0]
        if not fname.endswith('.png'): fname += '.png'
        local = os.path.join(IMG_DIR, fname)
        if os.path.exists(local):
            try:
                doc.add_picture(local, width=Inches(2.5))
                doc.add_paragraph(f'{"主图" if j==0 else "附图" if j < len(all_imgs)-len(var_imgs) else "变种图"} {j+1}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass

    doc.add_page_break()
    if (i+1) % 30 == 0: print(f'  DOCX: {i+1}/{rows}')

output_docx = os.path.join(OUT, 'Amazon回填表_产品清单.docx')
doc.save(output_docx)
print(f'\nDone! DOCX: {output_docx}')
print(f'Images: {IMG_DIR}')
